from pathlib import Path
import json
import shutil

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Sleep_Doomscrolling_Report.docx"
CHARTS = ROOT / "assets" / "charts"
S = json.loads((ROOT / "analysis_summary.json").read_text(encoding="utf-8"))

NAVY = "081120"
PANEL = "101D33"
BLUE = "335C9F"
CYAN = "1F9EB5"
VIOLET = "7259B8"
PINK = "B84B82"
GOLD = "C08B2A"
MINT = "248267"
RED = "B74357"
INK = "152238"
MUTED = "586A86"
PALE = "EEF4FB"
WHITE = "FFFFFF"
LINE = "CBD8E8"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=150, bottom=110, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa))); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd"); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent)); tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths_dxa:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(width)); grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx])); tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, name="Aptos", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None: run.font.size = Pt(size)
    if color: run.font.color.rgb = rgb(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(10.5); normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, VIOLET, 8, 4),
    ]:
        st = doc.styles[name]
        st.font.name = "Aptos Display"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = rgb(color)
        st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    cap = doc.styles["Caption"]
    cap.font.name = "Aptos"; cap.font.size = Pt(8.5); cap.font.italic = True; cap.font.color.rgb = rgb(MUTED)
    cap.paragraph_format.space_before = Pt(4); cap.paragraph_format.space_after = Pt(8)


def add_running_furniture(section, title="SLEEP × DOOMSCROLLING"):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    set_font(r, size=8, color=MUTED, bold=True)
    r.font.all_caps = True

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Night Signals Report  •  2026")
    set_font(r, size=8, color=MUTED)


def add_kicker(doc, text, color=CYAN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_font(r, size=8.5, color=color, bold=True)
    r.font.all_caps = True
    return p


def add_section_title(doc, icon, title, subtitle=None, color=BLUE):
    bar = doc.add_table(rows=1, cols=2)
    bar.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(bar, [650, 8710], indent=120)
    set_table_borders(bar, color=WHITE, size=0)
    shade(bar.cell(0,0), color); shade(bar.cell(0,1), PALE)
    bar.cell(0,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = bar.cell(0,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(icon), name="Segoe UI Symbol", size=17, color=WHITE, bold=True)
    p = bar.cell(0,1).paragraphs[0]
    set_font(p.add_run(title), name="Aptos Display", size=17, color=INK, bold=True)
    if subtitle:
        p2 = bar.cell(0,1).add_paragraph()
        p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(0)
        set_font(p2.add_run(subtitle), size=9, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def add_stat_strip(doc, stats, fill=PANEL):
    table = doc.add_table(rows=1, cols=len(stats))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [9360 // len(stats)] * len(stats)
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths, indent=120)
    set_table_borders(table, color=fill, size=0)
    for cell, (value, label) in zip(table.rows[0].cells, stats):
        shade(cell, fill); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(value), name="Aptos Display", size=19, color=WHITE, bold=True)
        p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        set_font(p2.add_run(label.upper()), size=7.5, color="B9CAE5", bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, label, text, color=CYAN):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [320, 9040], indent=120)
    set_table_borders(table, color=PALE, size=0)
    shade(table.cell(0,0), color); shade(table.cell(0,1), PALE)
    p = table.cell(0,1).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(label.upper()), size=8, color=color, bold=True)
    p2 = table.cell(0,1).add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run(text), name="Aptos Display", size=12.5, color=INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, filename, number, caption, takeaway, width=6.45):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(); r.add_picture(str(CHARTS / filename), width=Inches(width))
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.keep_with_next = True
    set_font(cp.add_run(f"Figure {number}. "), size=8.5, color=MUTED, bold=True)
    set_font(cp.add_run(caption), size=8.5, color=MUTED, italic=True)
    add_callout(doc, "Read this", takeaway, color=CYAN)


def add_story_map(doc):
    items = [
        ("01", "CORE", "Exposure → latency → wakeups → debt"),
        ("02", "MIND", "Negative news compounds mental load"),
        ("03", "SHIELD", "Routine and environment appear protective"),
        ("04", "EDGE", "Exceptions show risk is not destiny"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2340, 2340, 2340, 2340], indent=120)
    set_table_borders(table, color=LINE, size=6)
    for cell, (num, label, body) in zip(table.rows[0].cells, items):
        shade(cell, "F5F8FC")
        p = cell.paragraphs[0]
        set_font(p.add_run(num), name="Aptos Display", size=16, color=CYAN, bold=True)
        p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
        set_font(p2.add_run(label), size=8, color=BLUE, bold=True)
        p3 = cell.add_paragraph(); p3.paragraph_format.space_after = Pt(0)
        set_font(p3.add_run(body), size=8.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items, color=BLUE):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(.5)
        p.paragraph_format.first_line_indent = Inches(-.25)
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(item), size=10.5, color=INK)


def add_persona_card(doc, title, tag, body, metrics, color):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent=120)
    set_table_borders(table, color=color, size=10)
    cell = table.cell(0,0); shade(cell, PALE)
    p = cell.paragraphs[0]
    set_font(p.add_run(title), name="Aptos Display", size=15, color=color, bold=True)
    set_font(p.add_run(f"   {tag}"), size=8.5, color=MUTED, bold=True)
    p2 = cell.add_paragraph(); set_font(p2.add_run(body), size=10, color=INK)
    p3 = cell.add_paragraph(); p3.paragraph_format.space_after = Pt(0)
    for i, (label, value) in enumerate(metrics):
        if i: set_font(p3.add_run("    |    "), size=9, color=LINE)
        set_font(p3.add_run(f"{value} "), size=11, color=color, bold=True)
        set_font(p3.add_run(label), size=8.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


doc = Document()
style_doc(doc)
add_running_furniture(doc.sections[0])

# Cover page - editorial cover with a strong night panel.
cover = doc.add_table(rows=1, cols=1)
cover.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(cover, [9360], indent=120)
set_table_borders(cover, color=NAVY, size=0)
shade(cover.cell(0,0), NAVY)
cell = cover.cell(0,0)
set_cell_margins(cell, top=520, start=520, bottom=520, end=520)
p = cell.paragraphs[0]
set_font(p.add_run("NIGHT SIGNALS / 2026"), size=9, color="8BD9E8", bold=True)
p.paragraph_format.space_after = Pt(24)
p = cell.add_paragraph()
set_font(p.add_run("Sleep, Doomscrolling\nand the Night"), name="Aptos Display", size=31, color=WHITE, bold=True)
p.paragraph_format.space_after = Pt(12)
p = cell.add_paragraph()
set_font(p.add_run("How late-night scrolling, negative news, and protective routines relate to sleep quality across 1,000 respondents"), size=13, color="C9D8EF")
p.paragraph_format.space_after = Pt(30)
p = cell.add_paragraph()
set_font(p.add_run("DATA ANALYSIS REPORT"), size=8.5, color="8BD9E8", bold=True)
set_font(p.add_run("   •   29 analytic variables   •   11 report figures"), size=8.5, color="AFC1DE")

doc.add_paragraph().paragraph_format.space_after = Pt(24)
add_stat_strip(doc, [
    ("47.6%", "Doomscrollers"),
    ("+10.6 min", "Latency gap"),
    ("11 / 204", "Heavy-scroller exceptions"),
], fill=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Prepared from sleep_doomscrolling_habits.csv  |  Observational synthetic-data analysis"), size=8.5, color=MUTED, italic=True)
doc.add_page_break()

# Executive summary
add_section_title(doc, "☾", "Executive summary", "The answer first: the strongest signal is the bedtime disruption chain", color=BLUE)
add_body(doc, "Doomscrolling is consistently associated with a harder night in this dataset. Respondents labeled doomscrollers take longer to fall asleep, wake more often, carry more weekly sleep debt, and sleep slightly less. The clearest continuous marker is bedtime screen time, which rises alongside latency and disruption.")
add_body(doc, "The relationship is not destiny. A small set of heavy scrollers still report good sleep, and they more often pair their exposure with restorative routines, longer realized sleep, and lower debt. That counter-pattern is important because it shifts the practical question from ‘Who scrolls?’ to ‘What surrounds the scrolling?’")
add_stat_strip(doc, [
    (f"{S['latency_gap_minutes']:.1f} min", "Longer latency"),
    (f"{S['sleep_debt_gap_hours']:.1f} h", "More weekly debt"),
    (f"{abs(S['sleep_duration_gap_hours'])*60:.0f} min", "Less sleep nightly"),
], fill=PANEL)
add_bullets(doc, [
    "Mental load compounds behavioral exposure: doomscrolling plus negative-news consumption aligns with the highest anxiety, stress, and fatigue averages.",
    "Reading, meditation/journaling, a phone outside the bedroom, and regular exercise show more protective patterns than night mode alone.",
    "Behavioral variables predict sleep-quality category more strongly than most demographics, although subgroup context still matters for targeting.",
    "The data appear synthetic or simulation-assisted; findings are descriptive, non-causal, and should not be generalized as population prevalence.",
])
add_story_map(doc)
doc.add_page_break()

# Hero chart
add_kicker(doc, "Hero finding", CYAN)
doc.add_heading("More scrolling shifts the night toward poorer sleep", level=1)
add_body(doc, "Nightly doomscroll load combines session count with average session duration. Poor sleepers sit noticeably farther to the right, but the scatter also reveals overlap and exceptions—useful evidence that risk accumulates without becoming deterministic.")
add_figure(doc, "02_hero_doomscroll_sleep.png", 1,
           "Nightly doomscroll load by reported sleep-quality category.",
           "Poor sleepers carry the highest median nightly doomscroll load; good sleepers cluster lower, with meaningful overlap.")

# Core relationship
add_section_title(doc, "▣", "Core relationship", "Q1–2  |  Sleep quality, latency, wakeups, duration, and debt", color=VIOLET)
add_body(doc, "Across five outcomes, the doomscroller label points in the same direction: longer latency, more wakeups, more debt, shorter sleep, and greater fatigue. Consistency across outcomes strengthens the descriptive story, even though shared causes—stress, habit intensity, or chronotype—may explain part of every gap.")
add_figure(doc, "03_doomscroller_comparison.png", 2,
           "Mean sleep outcomes for doomscrollers and non-doomscrollers; error bars show approximate 95% confidence intervals.",
           "Doomscrollers average 10.6 extra minutes of latency and 1.5 additional hours of weekly debt, plus more wakeups and fatigue.")
add_callout(doc, "Pull quote", "The night does not break in one place—it stretches, fragments, and leaves a daytime trace.", color=VIOLET)

add_kicker(doc, "Dose response", PINK)
doc.add_heading("Bedtime exposure matters beyond the binary label", level=1)
add_body(doc, "Quartiles reveal a graded pattern: the higher the bedtime screen exposure, the longer respondents take to fall asleep and the more sleep debt they carry. The same direction appears in wakeups and sleep duration, suggesting that the binary doomscroller label compresses a richer dose-response relationship.")
add_figure(doc, "04_dose_response.png", 3,
           "Average outcomes across quartiles of bedtime screen time.",
           "Latency and debt climb across screen-time quartiles while sleep duration declines and wakeups rise.")

# Mental health
add_section_title(doc, "◉", "Mental health angle", "Q3–4  |  Anxiety, stress, fatigue, and negative-news consumption", color=PINK)
add_body(doc, "Doomscrolling is as much about content and emotional activation as it is about minutes. When negative-news consumption accompanies doomscrolling, average anxiety, stress, and fatigue are highest. This may reflect a feedback loop: distress can fuel checking behavior, and repeated exposure can prolong arousal.")
add_figure(doc, "05_mental_health.png", 4,
           "Average anxiety, stress, and fatigue by doomscrolling and negative-news status.",
           "The dual-exposure group—doomscrolling plus negative news—shows the heaviest mental-wellbeing burden.")
add_callout(doc, "Interpret carefully", "Direction is unresolved: the dataset cannot tell whether distress drives scrolling, scrolling drives distress, or both.", color=PINK)

# Protective habits
add_section_title(doc, "◇", "Protective habits", "Q5–6  |  What appears to soften the pattern?", color=MINT)
add_body(doc, "The most useful protective signals are structural: what the respondent does before bed, where the phone sleeps, and whether daytime movement is protected. Night mode is convenient, but its average profile is weaker than the profile of a genuinely different bedtime routine.")
add_figure(doc, "06_protective_habits.png", 5,
           "Good-sleep share and sleep latency for selected habits and routines.",
           "Reading and meditation/journaling stand out; a social-scrolling routine performs worst, and night mode alone is not a strong substitute for routine change.")
add_figure(doc, "07_exercise_gradient.png", 6,
           "Good-sleep share and daytime fatigue across exercise bands.",
           "Exercise looks modestly protective, but the pattern is not perfectly monotonic and should not be treated as a cure-all.")

# Demographics
add_section_title(doc, "◎", "Demographic context", "Q7–8  |  Age buckets, occupation, country, and descriptive gender context", color=GOLD)
add_body(doc, "Age buckets replace raw age wherever they make the story easier to read: Teens (15–19), 20s (20–29), and 30s+. Age–occupation cells reveal pockets of elevated doomscrolling, while country rates vary across samples of very different sizes. These breakdowns are useful for segmentation, not ranking.")
add_figure(doc, "08_demographics.png", 7,
           "Doomscrolling by age bucket and occupation, plus country-level doomscrolling and poor-sleep shares.",
           "Younger/student pockets are often elevated, but small cells and uneven country samples make the differences descriptive rather than definitive.")
add_callout(doc, "Gender", "Gender differences are modest in this sample and are presented descriptively only; behavior remains the more actionable signal.", color=GOLD)

# Exceptions
add_section_title(doc, "↗", "The Exceptions", "Counter-narrative  |  Heavy scrollers who still report good sleep", color=CYAN)
add_body(doc, f"Heavy doomscrollers are defined here as self-identified doomscrollers at or above {S['heavy_threshold']:.0f} minutes of bedtime screen time—the top quartile threshold. There are {S['heavy_n']} such respondents, yet only {S['exception_n']} report good sleep. The group is small enough to demand caution, but distinct enough to challenge a simplistic ‘scrolling equals poor sleep’ conclusion.")
add_figure(doc, "10_exceptions.png", 8,
           "Standardized differences between good-sleep exceptions and other heavy scrollers.",
           "The exceptions realize more sleep and less debt/latency, with restorative bedtime routines appearing more often.")
add_callout(doc, "Counter-narrative", "Exposure raises risk; routine and realized sleep appear to shape whether that risk becomes an outcome.", color=CYAN)

# Personas
add_section_title(doc, "◌", "Personas", "Three transparent respondent profiles for product and behavior design", color=VIOLET)
add_body(doc, "These personas are manually defined combinations, not diagnoses or opaque clusters. They make the evidence easier to act on by separating exposure-heavy, emotion-heavy, and routine-protected patterns.")
add_persona_card(doc, "The Night Scroller", "n=204", "High bedtime exposure and a self-identified doomscrolling pattern. The first intervention is environmental: reduce easy access, add a stop cue, and move the phone away from the pillow.", [("bedtime screen", "Top quartile"), ("primary lever", "Friction"), ("sleep risk", "High")], PINK)
add_persona_card(doc, "The Anxious News Seeker", "n=51", "Negative-news consumption with high anxiety and stress. The first intervention is content-aware: limit news windows, remove alerts, and replace late checking with a decompression ritual.", [("emotional load", "High"), ("primary lever", "Content"), ("sleep risk", "Elevated")], GOLD)
add_persona_card(doc, "The Disciplined Sleeper", "n=130", "Lower bedtime exposure paired with reading or meditation/journaling. The design opportunity is maintenance: protect the routine and make the good choice the default.", [("screen exposure", "Lower"), ("primary lever", "Routine"), ("sleep profile", "Protective")], MINT)
add_figure(doc, "11_personas.png", 9,
           "Relative persona profiles with observed metric values printed in each cell.",
           "Different mechanisms imply different interventions: friction for exposure, content boundaries for distress, and reinforcement for disciplined routines.")

# Synthesis
add_section_title(doc, "∑", "Synthesis", "Q9  |  Correlation structure and predictive ranking", color=BLUE)
add_body(doc, "The correlation structure resembles a bedtime-disruption chain. Exposure relates to latency; latency, wakeups, and shorter sleep accumulate into debt and fatigue. The same coherence that makes the story understandable also heightens the synthetic-data caveat: multiple variables may have been jointly generated around one latent disruption score.")
add_figure(doc, "12_correlation_heatmap.png", 10,
           "Pearson correlations among selected behavioral, sleep, and wellbeing measures.",
           "Bedtime screen time, doomscroll sessions, and latency form the strongest exposure cluster; sleep duration moves opposite sleep debt.")

add_kicker(doc, "Predictive view", BLUE)
doc.add_heading("The behavior label leads; the night’s mechanics follow", level=1)
add_body(doc, f"A random-forest classifier evaluated with five-fold cross-validation reaches {S['cv_balanced_accuracy']:.1%} balanced accuracy against a {S['baseline_accuracy']:.1%} majority-class baseline. Permutation importance is computed out of fold and grouped at the original feature level. The outcome-adjacent sleep_quality_score is excluded to avoid construct-overlap leakage.")
add_figure(doc, "13_feature_importance.png", 11,
           "Cross-validated permutation importance for predicting sleep-quality category.",
           "The doomscroller label leads, followed by wakeups, sleep duration, sleep debt, and latency; correlated exposure variables divide their shared signal.")
add_callout(doc, "Model boundary", "Prediction ranks signals; it does not identify causes, prescribe treatment, or validate the survey against real-world prevalence.", color=BLUE)

# Conclusion
add_section_title(doc, "✓", "Conclusion & practical takeaways", "A sequence for testing, not a medical prescription", color=MINT)
add_body(doc, "The most defensible conclusion is not that every minute of scrolling destroys sleep. It is that a repeatable bedtime-disruption pattern—longer exposure, more checking, delayed sleep, more wakeups, and accumulated debt—appears throughout the data. Practical action should therefore target the chain rather than moralize the behavior.")
add_bullets(doc, [
    "Create friction before bedtime: charge the phone outside the bedroom or beyond arm’s reach.",
    "Replace, do not merely remove: choose a fixed reading, meditation, or journaling routine for the final 20–30 minutes.",
    "Treat negative news as a separate lever: use scheduled news windows and disable late alerts.",
    "Protect daytime movement: aim for a consistent exercise habit without expecting it to offset severe nighttime disruption.",
    "Track latency, wakeups, and weekly debt—not only total screen time—to see whether an intervention changes sleep mechanics.",
    "Test by persona: exposure-heavy users need friction, distress-heavy users need content boundaries, and disciplined sleepers need maintenance cues.",
])
add_callout(doc, "Best next test", "A two-week phone-outside-bedroom experiment with a fixed replacement routine and daily latency tracking.", color=MINT)

doc.add_heading("Limitations", level=1)
add_bullets(doc, [
    "The dataset appears synthetic or simulation-assisted: target classes are unusually balanced, relationships are unusually clean, and valid sleep-quality scores are heavily concentrated at 5/5.",
    "The design is cross-sectional and self-reported; temporal direction and causal effects cannot be established.",
    "Median imputation and an Unknown category preserve all 1,000 records but can attenuate or redistribute associations.",
    "Some subgroup cells—especially the good-sleep heavy-scroller exception group—are small and should be treated as exploratory.",
    "Sleep-quality category is a coarse target. The model excludes sleep_quality_score to reduce direct construct overlap, but other sleep variables remain outcome-adjacent by design.",
    "Country and gender comparisons are descriptive only; unequal sample sizes and unmeasured social context prevent broad generalization.",
])

doc.add_heading("Method in brief", level=2)
add_body(doc, "The workflow validated IDs, dtypes, missingness, ranges, and potential formulaic relationships; created age buckets; compared groups and quartiles; defined exception and persona rules; visualized correlations; and evaluated a random-forest classifier with stratified five-fold cross-validation. Full executable detail, code, outputs, and figure takeaways are in sleep_doomscrolling_analysis.ipynb.")

doc.save(OUT)
public_downloads = ROOT / "public" / "downloads"
public_downloads.mkdir(parents=True, exist_ok=True)
shutil.copy2(OUT, public_downloads / OUT.name)
print(f"Saved {OUT}")
